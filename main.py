from fastapi import FastAPI, Request, Form, HTTPException
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, Response
from sse_starlette.sse import EventSourceResponse
import os
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docxtpl import DocxTemplate
from datetime import datetime, timedelta
import random
from PyPDF2 import PdfMerger
from docx2pdf import convert
import re
import asyncio
import json
from typing import AsyncGenerator, Dict
from pypdf import PdfWriter, PdfReader
import shutil
import urllib.parse  # 添加这个导入
from pathlib import Path
from docx.enum.text import WD_BREAK  # 添加这行导入
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from concurrent.futures import ThreadPoolExecutor
import traceback  # 添加这个导入
import zipfile
import tempfile
from starlette.background import BackgroundTask

app = FastAPI()

# 挂载静态文件目录
app.mount("/static", StaticFiles(directory="static"), name="static")

# 设置模板目录
templates = Jinja2Templates(directory="templates")

# 获取在线数据
try:
    from get_new_data import get_new_data
    product_json, factory_list = get_new_data()
except ImportError as e:
    print("导入 get_new_data 失败:", e)
    product_json, factory_list = [], []


# 全局变量存储进度和路径
pdf_progress = {}
pdf_paths = {}
pdf_final_paths = {}
merge_progress: Dict[str, int] = {}
merge_paths: Dict[str, str] = {}

def deal_input_number(user_input):
    return_list = []
    user_input_list = user_input.strip().split(' ')
    for item in user_input_list:
        if '-' in item:
            first_num = int(float(item.split('-')[0]))
            second_num = int(float(item.split('-')[1]))
            for i in range(first_num, second_num + 1):
                return_list.append("{:03d}".format(i))
        else:
            return_list.append("{:03d}".format(int(float(item))))
    return return_list

def get_all_problem_files_name(company_name, user_input_num):
    problem_nums_list = deal_input_number(user_input_num)
    all_files = os.listdir(f"{company_name}")
    all_problem_files_name = []
    for problem_num in problem_nums_list:
        for file in all_files:
            if problem_num == file.split('-')[1][-3:]:
                all_problem_files_name.append(f"{company_name}/{file}")
    return all_problem_files_name

def change_one_brand(problem_file_name, alert_factory, alert_type):
    factory_name = [item['fullname'] for item in product_json if item['name'] == alert_factory][0]
    dongzuozhi = [item['alarm'] for item in product_json if item['name'] == alert_factory][0]

    document = Document(problem_file_name)
    document.paragraphs[16].runs[6].text = "    " + alert_type
    document.paragraphs[22].runs[9].text = factory_name
    document.paragraphs[62].runs[3].text = str(dongzuozhi) + "%LEL"

    document.save(problem_file_name)

def change_all_brand(company_name, product_company, product_model, product_num_list):
    try:
        problem_files_name = get_all_problem_files_name(company_name, product_num_list)
        if not problem_files_name:
            return {"status": "error", "message": "未找到需要修改的文件"}
            
        for problem_file_name in problem_files_name:
            change_one_brand(problem_file_name, product_company, product_model.split('（')[0])
            
        return {"status": "success", "message": "品牌已经修改完成，请到指定的公司文件夹查看"}
        
    except Exception as e:
        return {"status": "error", "message": str(e)}

def change_problem_file(problem_filename):
    document = Document(problem_filename)

    document.paragraphs[60].runs[2].text = "异常"  # 更改报警功能的值
    document.paragraphs[62].runs[3].text = "/    "  # 更改报警动作值
    document.paragraphs[62].runs[4].text = ""
    document.paragraphs[65].runs[3].text = "/"  # 更改重复性
    document.paragraphs[65].runs[4].text = ""
    document.paragraphs[66].runs[3].text = "  /   "  # 更改响应时间
    document.paragraphs[66].runs[4].text = ""

    tables = document.tables
    table = tables[1]
    table.cell(1, 1).text = "/"
    table.cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(1, 1).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(1, 2).text = "/"
    table.cell(1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(1, 2).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(1, 3).text = "/"
    table.cell(1, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(1, 3).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(2, 1).text = "/"
    table.cell(2, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(2, 1).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(2, 2).text = "/"
    table.cell(2, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(2, 2).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(2, 3).text = "/"
    table.cell(2, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(2, 3).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(3, 1).text = "/"
    table.cell(3, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(3, 1).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(3, 2).text = "/"
    table.cell(3, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(3, 2).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(3, 3).text = "/"
    table.cell(3, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(3, 3).paragraphs[0].runs[0].font.size = Pt(12)

    document.save(problem_filename)

def find_text_in_word(text):
    filename = "test.docx"
    document = Document(filename)
    target_text = text
    # 遍历文档中的每个段落
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        # 遍历段落中的每个运行
        for run_index, run in enumerate(paragraph.runs):
            # 判断运行的文本是否与目标文本匹配
            if target_text in run.text:
                # 返回匹配到的段落序号和运行序号
                result = f"段落序号：{paragraph_index}，运行序号：{run_index}"
                return paragraph_index, run_index
    return None

@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    return templates.TemplateResponse(
        "index.html", 
        {"request": request}
    )

@app.get("/problem", response_class=HTMLResponse)
async def problem_page(request: Request):
    return templates.TemplateResponse(
        "problem.html", 
        {"request": request}
    )

@app.get("/brand", response_class=HTMLResponse)
async def brand_page(request: Request):

    
    # 从product_json中提取品牌和型号数据
    brands_data = {}
    
    for product in product_json:
        brand_name = product['name']  # 品牌简称（用作value）
        full_name = product['fullname']  # 公司名称（用作显示）
        # 暂时把完整的产品数据传到前端，方便调试
        brands_data[brand_name] = product
    
    
    return templates.TemplateResponse(
        "brand.html", 
        {
            "request": request,
            "brands_data": brands_data
        }
    )

@app.post("/process/problem")
async def process_problem(
    company_name: str = Form(...),
    probe_numbers: str = Form(...)
):
    try:
        if not os.path.exists(company_name):
            return {"status": "error", "message": f"找不到公司文件夹: {company_name}"}
        
        problem_files_name = get_all_problem_files_name(company_name, probe_numbers)
        if not problem_files_name:
            return {"status": "error", "message": "未找到匹配的文件"}
        
        for problem_file_name in problem_files_name:
            change_problem_file(problem_file_name)
        
        return {"status": "success", "message": f"成功处理了 {len(problem_files_name)} 个故障探头文件"}
    
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/process/brand")
async def process_brand(
    company_name: str = Form(...),
    probe_numbers: str = Form(...),
    product_company: str = Form(...),
    product_model: str = Form(...)
):
    try:
        if not os.path.exists(company_name):
            return {"status": "error", "message": f"找不到公司文件夹: {company_name}"}
        
        return change_all_brand(company_name, product_company, product_model, probe_numbers)
    
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.get("/generate", response_class=HTMLResponse)
async def generate_page(request: Request):
    return templates.TemplateResponse(
        "generate.html", 
        {"request": request}
    )

def parse_probe_numbers(probe_numbers_str):
    result = []
    parts = probe_numbers_str.split()
    for part in parts:
        if '-' in part:
            start, end = map(int, part.split('-'))
            result.extend(range(start, end + 1))
        else:
            result.append(int(part))
    return sorted(result)

def generate_certificate(template_path, output_path, probe_num, test_date, report_date):
    doc = DocxTemplate(template_path)
    context = {
        'probe_num': probe_num,
        'test_date': test_date,
        'report_date': report_date
    }
    doc.render(context)
    doc.save(output_path)

@app.post("/process/generate")
async def process_generate(
    company_name: str = Form(...),
    all_nums: int = Form(...),
    date: str = Form(...),
    temperature: float = Form(...),
    humidity: int = Form(...),
    sections: str = Form(...),
    sections_num: str = Form(...),
    start_num: int = Form(...)
):
    try:
        # 简化数据处理
        sections_list = sections.strip().split()
        sections_num_list = [int(x) for x in sections_num.strip().split()]
        formatted_date = date.replace("-", "")


        # 基本验证
        if not sections_list or not sections_num_list:
            raise ValueError("区域和数量不能为空")

        if len(sections_list) != len(sections_num_list):
            raise ValueError(
                f"区域数量与探头数量不匹配\n"
                f"区域：{', '.join(sections_list)}\n"
                f"数量：{', '.join(map(str, sections_num_list))}"
            )

        total_probes = sum(sections_num_list)
        if total_probes != all_nums:
            raise ValueError(
                f"探头总数与各区域数量之和不匹配\n"
                f"总数：{all_nums}\n"
                f"各区域之和：{total_probes}"
            )

        # 构建输入数据
        user_input = {
            "company_name": company_name,
            "all_nums": all_nums,
            "date": formatted_date,
            "temperature": float(temperature),
            "humidity": int(humidity),
            "sections": sections_list,
            "sections_num": sections_num_list,
            "start_num": start_num
        }

        # 创建文件夹
        folder_name = f"{company_name}{formatted_date}"
        os.makedirs(folder_name, exist_ok=True)

        # 生成证书
        if write_save_all_company(user_input):
            return {
                "status": "success",
                "message": f"证书已生成，请在「{folder_name}」文件夹中查看"
            }
        else:
            raise Exception("证书生成失败")

    except ValueError as ve:
        print(f"数据验证错误：{str(ve)}")
        return JSONResponse(
            status_code=422,
            content={"status": "error", "message": str(ve)}
        )
    except Exception as e:
        print(f"生成证书时出错：{str(e)}")
        return JSONResponse(
            status_code=422,
            content={"status": "error", "message": str(e)}
        )

# 添加新的辅助函数
def return_format_num(num):
    """根据序号返回三位的格式化编号"""
    if num < 10:
        return f"00{num}"
    elif num < 100:
        return f"0{num}"
    else:
        return f"{num}"

def get_file_num(date, num, start_num):
    """生成文件编号"""
    return f"ZJYX-{date}0{return_format_num(num + start_num)}"

def format_date(date_str):
    """格式化日期"""
    date = datetime.strptime(date_str, "%Y%m%d")
    formatted_date = date.strftime("%Y 年 %m  月 %d  日")
    next_year = date + timedelta(days=364)
    previous_day = next_year - timedelta(days=0)
    formatted_next_year_date = previous_day.strftime("%Y 年 %m  月 %d  日")
    return formatted_date, formatted_next_year_date

def create_all_alerts_num_list(sections, sections_num):
    """生成所有探头编号"""
    all_alerts_num = []
    
    # 打印调试信息
    print("创建探头编号列表:")
    print(f"sections: {sections}, type: {type(sections)}")
    print(f"sections_num: {sections_num}, type: {type(sections_num)}")
    
    # 确保输入是列表
    sections = list(sections)
    sections_num = [int(x) for x in sections_num]
    
    print("转换后:")
    print(f"sections: {sections}, type: {type(sections)}")
    print(f"sections_num: {sections_num}, type: {type(sections_num)}")
    
    try:
        for i in range(len(sections)):
            section = sections[i]
            num_probes = sections_num[i]
            for j in range(num_probes):
                alert_num = f"{section}{return_format_num(j + 1)}"
                all_alerts_num.append(alert_num)
                
        print(f"生成的探头编号: {all_alerts_num}")
        return all_alerts_num
        
    except Exception as e:
        print(f"生成探头编号时出错: {str(e)}")
        raise

def all_company_to_save(user_input):
    """生成所有公司数据"""
    try:
        all_company_data = []
        
        # 打印输入数据
        print("\n处理公司数据:")
        print(f"输入数据: {user_input}")
        
        # 提取并转换数据
        company_name = str(user_input["company_name"])
        all_nums = int(user_input["all_nums"])
        date = str(user_input["date"])
        temperature = float(user_input["temperature"])
        humidity = int(user_input["humidity"])
        sections = list(user_input["sections"])
        sections_num = list(user_input["sections_num"])
        start_num = int(user_input["start_num"])
        
        # 生成探头编号
        all_alerts_num = create_all_alerts_num_list(sections, sections_num)
        print(f"生成的探头编号列表: {all_alerts_num}")
        
        # 生成日期
        date_now, date_next = format_date(date)
        
        # 生成每个证书的数据
        for i in range(all_nums):
            new_company = {
                "file_num": get_file_num(date, i, start_num),
                "company_name": company_name,
                "alert_type": "AEC2332",
                "alert_factory": "成都安可信电子股份有限公司",
                "dongzuozhi": "ankexindongzuo",
                "alert_num": all_alerts_num[i],
                "date_now": date_now,
                "date_next": date_next,
                "temperature": temperature,
                "humidity": humidity,
                "random_chongfu": round(random.uniform(0.5, 2.0), 1),
                "action_time": random.randint(7, 25)
            }
            all_company_data.append(new_company)
            
        print(f"生成的证书数据数量: {len(all_company_data)}")
        return all_company_data
        
    except Exception as e:
        print(f"生成公司数据时出错: {str(e)}")
        raise

def change_table_cell(document):
    """修改表格数据并调整格式"""
    try:
        # 读取参考表格
        document_table = Document("table_refer.docx")
        all_tables = list(document_table.tables)  # 转换为列表
        
        # 随机选择一个表格
        random_table_num = random.randint(0, len(all_tables) - 1)
        random_table = all_tables[random_table_num]
        
        # 获取随机值
        random_value_1 = random_table.cell(1, 1).text
        random_value_2 = random_table.cell(2, 1).text
        random_value_3 = random_table.cell(3, 1).text
        random_diff_1 = random_table.cell(1, 2).text
        random_diff_2 = random_table.cell(2, 2).text
        random_diff_3 = random_table.cell(3, 2).text
        
        # 获取目标文档的表格
        tables = list(document.tables)  # 转换为列表
        if len(tables) < 2:
            raise ValueError("目标文档中没有足够的表格")
            
        table = tables[1]
        
        # 设置单元格内容和格式
        cells_to_update = [
            (1, 1, random_value_1),
            (2, 1, random_value_2),
            (3, 1, random_value_3),
            (1, 2, random_diff_1),
            (2, 2, random_diff_2),
            (3, 2, random_diff_3)
        ]
        
        for row, col, value in cells_to_update:
            cell = table.cell(row, col)
            cell.text = value
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(12)
        
        return document
        
    except Exception as e:
        print(f"修改表格时出错: {str(e)}")
        raise

def write_save_all_company(user_input):
    """保存所有公司证书"""
    try:
        print("\n开始生成证书:")
        
        # 创建公司文件夹
        folder_name = f"{user_input['company_name']}{user_input['date']}"
        os.makedirs(folder_name, exist_ok=True)
        print(f"创建文件夹: {folder_name}")

        # 生成所有证书数据
        all_company_data = all_company_to_save(user_input)
        print(f"生成证书数据完成，共 {len(all_company_data)} 条")

        # 生成每个证书
        for i, company in enumerate(all_company_data):
            doc = DocxTemplate("model.docx")
            doc.render(company)
            change_table_cell(doc)
            file_path = os.path.join(folder_name, f"{company['file_num']}-{company['alert_num']}.docx")
            doc.save(file_path)
            print(f"生成第 {i+1} 个证书: {file_path}")

        return True

    except Exception as e:
        print(f"生成证书时出错: {str(e)}")
        raise

@app.get("/pdf")
async def pdf_page(request: Request):
    return templates.TemplateResponse("pdf.html", {"request": request})

@app.post("/process/pdf")
async def process_pdf(company_folder: str = Form(...)):
    try:
        current_dir = os.getcwd()
        company_path = os.path.join(current_dir, company_folder)

        # 检查公司文件夹是否存在
        if not os.path.exists(company_path):
            return JSONResponse(
                status_code=422,
                content={
                    "status": "error",
                    "message": f"找不到公司文件夹：{company_folder}"
                }
            )

        # 检查Word文件
        docx_files = [
            f for f in os.listdir(company_path) 
            if f.lower().endswith('.docx') and not f.startswith('~$')
        ]

        if not docx_files:
            return JSONResponse(
                status_code=422,
                content={
                    "status": "error",
                    "message": "在文件夹中未找到Word文件"
                }
            )

        # 生成请求ID并存储路径
        request_id = str(abs(hash(company_path)))
        pdf_progress[request_id] = 0
        pdf_paths[request_id] = company_path
        
        # 启动处理任务
        asyncio.create_task(process_pdf_async(request_id))
        
        return JSONResponse(
            status_code=200,
            content={
                "status": "success",
                "request_id": request_id,
                "message": f"开始处理，共找到 {len(docx_files)} 个Word文件"
            }
        )
    except Exception as e:
        return JSONResponse(
            status_code=422,
            content={"status": "error", "message": str(e)}
        )

async def process_pdf_async(request_id: str):
    try:
        company_path = pdf_paths.get(request_id)
        if not company_path:
            raise ValueError("找不到公司文件夹路径")

        # 创建PDF证书文件夹
        pdf_folder = os.path.join(company_path, "PDF证书")
        if os.path.exists(pdf_folder):
            shutil.rmtree(pdf_folder)
        os.makedirs(pdf_folder)
        print(f"创建PDF文件夹: {pdf_folder}")
        
        # 获取所有Word文件并按名称排序
        docx_files = [
            f for f in os.listdir(company_path) 
            if f.lower().endswith('.docx') and not f.startswith('~$')
        ]
        docx_files.sort()
        total_files = len(docx_files)
        
        if not docx_files:
            raise ValueError("没有找到Word文件")
        
        print(f"开始转换，共 {total_files} 个文件")
        
        # 转换每个Word文件为PDF
        converted_files = []
        for i, docx_file in enumerate(docx_files, 1):
            try:
                input_path = os.path.join(company_path, docx_file)
                output_name = os.path.splitext(docx_file)[0] + '.pdf'
                output_path = os.path.join(pdf_folder, output_name)
                print(f"转换第 {i}/{total_files} 个文件: {docx_file}")
                print(f"输入路径: {input_path}")
                print(f"输出路径: {output_path}")
                
                # 检查输入文件
                if not os.path.exists(input_path):
                    print(f"输入文件不存在: {input_path}")
                    continue
                
                try:
                    # 使用线程池执行转换
                    loop = asyncio.get_event_loop()
                    await loop.run_in_executor(
                        None,
                        lambda: convert(input_path, output_path, keep_active=False)
                    )
                except Exception as conv_error:
                    print(f"转换过程出错: {str(conv_error)}")
                    print("详细错误信息:")
                    print(traceback.format_exc())
                    continue
                
                # 检查转换结果
                if os.path.exists(output_path):
                    if os.path.getsize(output_path) > 0:
                        converted_files.append(output_path)
                        print(f"转换成功: {output_path}")
                    else:
                        print(f"转换后的文件大小为0: {output_path}")
                else:
                    print(f"转换后的文件不存在: {output_path}")
                
                pdf_progress[request_id] = int((i / total_files) * 60)
                await asyncio.sleep(0.1)
                
            except Exception as e:
                print(f"处理文件时出错: {str(e)}")
                print("详细错误信息:")
                print(traceback.format_exc())
                continue
        
        # 检查转换结果
        print(f"成功转换的文件数: {len(converted_files)}")
        if not converted_files:
            raise ValueError("没有成功转换的PDF文件")
        
        # 合并PDF文件
        print("开始合并PDF文件")
        merger = PdfWriter()
        
        for i, pdf_path in enumerate(converted_files, 1):
            print(f"合并文件 {i}/{len(converted_files)}: {pdf_path}")
            try:
                pdf = PdfReader(pdf_path)
                for page in pdf.pages:
                    merger.add_page(page)
                
                progress = 60 + int((i / len(converted_files)) * 40)
                pdf_progress[request_id] = progress
                await asyncio.sleep(0.1)
            except Exception as e:
                print(f"合并PDF时出错: {str(e)}")
                print("详细错误信息:")
                print(traceback.format_exc())
                continue
        
        # 保存合并后的PDF
        output_name = f"{os.path.basename(company_path)}合并证书.pdf"
        output_path = os.path.join(company_path, output_name)
        print(f"保存合并文件: {output_path}")
        
        try:
            with open(output_path, 'wb') as output_file:
                merger.write(output_file)
        except Exception as e:
            print(f"保存合并文件时出错: {str(e)}")
            print("详细错误信息:")
            print(traceback.format_exc())
            raise
        
        if os.path.getsize(output_path) == 0:
            raise ValueError("生成的PDF文件为空")
        
        # 保存文件路径到全局变量
        pdf_paths[request_id] = output_path
        pdf_progress[request_id] = 100
        print("合并完成")
        
    except Exception as e:
        print(f"处理出错: {str(e)}")
        print("详细错误信息:")
        print(traceback.format_exc())
        if request_id in pdf_progress:
            del pdf_progress[request_id]
        if request_id in pdf_paths:
            del pdf_paths[request_id]
        raise

@app.get("/process/pdf/progress/{request_id}")
async def pdf_progress_stream(request_id: str):
    try:
        if request_id not in pdf_progress:
            return JSONResponse(
                status_code=404,
                content={"status": "error", "message": "无效的请求ID"}
            )

        async def generate_progress():
            while pdf_progress.get(request_id, 0) < 100:
                if request_id in pdf_progress:
                    current_progress = pdf_progress[request_id]
                    status_message = (
                        "正在转换Word文件..." if current_progress <= 60 
                        else "正在合并PDF文件..."
                    )
                    yield {
                        "data": json.dumps({
                            "progress": current_progress / 100,
                            "status": "processing",
                            "message": f"{status_message} ({current_progress}%)"
                        })
                    }
                await asyncio.sleep(0.5)
            
            yield {
                "data": json.dumps({
                    "progress": 1.0,
                    "status": "completed",
                    "message": "文件已合并完成并保存到桌面！"
                })
            }
            
            # 只清理进度信息
            if request_id in pdf_progress:
                del pdf_progress[request_id]

        return EventSourceResponse(generate_progress())
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": str(e)}
        )

@app.get("/download/{request_id}")
async def download_pdf(request_id: str):
    try:
        pdf_path = pdf_paths.get(request_id)
        if not pdf_path:
            raise HTTPException(status_code=404, detail="文件不存在或已过期")
            
        return FileResponse(
            pdf_path,
            media_type='application/pdf',
            filename=os.path.basename(pdf_path)
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/merge")
async def merge_page(request: Request):
    return templates.TemplateResponse("merge.html", {"request": request})

@app.post("/process/merge")
async def process_merge(company_folder: str = Form(...)):
    try:
        current_dir = os.getcwd()
        company_path = os.path.join(current_dir, company_folder)

        # 检查公司文件夹是否存在
        if not os.path.exists(company_path):
            return JSONResponse(
                status_code=422,
                content={
                    "status": "error",
                    "message": f"找不到公司文件夹：{company_folder}"
                }
            )

        # 检查Word文件
        docx_files = [
            f for f in os.listdir(company_path) 
            if f.lower().endswith('.docx') and not f.startswith('~$')
        ]

        if not docx_files:
            return JSONResponse(
                status_code=422,
                content={
                    "status": "error",
                    "message": "在文件夹中未找到Word文件"
                }
            )

        # 生成请求ID并存储路径
        request_id = str(abs(hash(company_path)))
        merge_progress[request_id] = 0
        merge_paths[request_id] = company_path
        
        # 启动处理任务
        asyncio.create_task(process_merge_async(request_id))
        
        return JSONResponse(
            status_code=200,
            content={
                "status": "success",
                "request_id": request_id,
                "message": f"开始处理，共找到 {len(docx_files)} 个Word文件"
            }
        )
    except Exception as e:
        return JSONResponse(
            status_code=422,
            content={"status": "error", "message": str(e)}
        )

async def process_merge_async(request_id: str):
    try:
        company_path = merge_paths.get(request_id)
        if not company_path:
            raise ValueError("找不到公司文件夹路径")

        # 获取所有Word文件并按名称排序
        docx_files = [
            f for f in os.listdir(company_path) 
            if f.lower().endswith('.docx') and not f.startswith('~$')
        ]
        docx_files.sort()
        
        if not docx_files:
            raise ValueError("没有找到Word文件")
            
        print(f"开始合并，共 {len(docx_files)} 个文件")
        
        # 创建输出文件名和路径
        output_name = f"{os.path.basename(company_path)}合并证书.docx"
        output_path = os.path.join(company_path, output_name)
        
        # 创建新的空白文档
        merged_doc = Document()
        
        # 处理所有文档
        total_files = len(docx_files)
        for i, docx_file in enumerate(docx_files):
            print(f"正在处理第 {i+1}/{total_files} 个文件: {docx_file}")
            file_path = os.path.join(company_path, docx_file)
            doc = Document(file_path)
            
            # 复制当前文档的所有内容
            for element in doc.element.body:
                if element.tag.endswith(('p', 'tbl', 'sectPr')):
                    merged_doc.element.body.append(element)
            
            # 如果不是最后一个文件，添加分页符
            if i < total_files - 1:
                print(f"在文件 {docx_file} 后添加分页符")
                paragraph = merged_doc.add_paragraph()
                run = paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
            
            # 更新进度
            merge_progress[request_id] = int(((i + 1) / total_files) * 100)
            await asyncio.sleep(0.1)
        
        print(f"保存合并文件: {output_path}")
        merged_doc.save(output_path)
        
        # 检查生成的文件
        if os.path.getsize(output_path) == 0:
            raise ValueError("生成的Word文件为空")
        
        # 保存文件路径到全局变量
        merge_paths[request_id] = output_path
        merge_progress[request_id] = 100
        print("合并完成")
        
    except Exception as e:
        print(f"处理出错: {str(e)}")
        if request_id in merge_progress:
            del merge_progress[request_id]
        if request_id in merge_paths:
            del merge_paths[request_id]
        raise

@app.get("/process/merge/progress/{request_id}")
async def merge_progress_stream(request_id: str):
    try:
        if request_id not in merge_progress:
            return JSONResponse(
                status_code=404,
                content={"status": "error", "message": "无效的请求ID"}
            )

        async def generate_progress():
            while merge_progress.get(request_id, 0) < 100:
                if request_id in merge_progress:
                    current_progress = merge_progress[request_id]
                    yield {
                        "data": json.dumps({
                            "progress": current_progress / 100,
                            "status": "processing",
                            "message": f"正在合并Word文件... ({current_progress}%)"
                        })
                    }
                await asyncio.sleep(0.5)
            
            yield {
                "data": json.dumps({
                    "progress": 1.0,
                    "status": "completed",
                    "message": "Word文件已合并完成！"
                })
            }
            
            if request_id in merge_progress:
                del merge_progress[request_id]

        return EventSourceResponse(generate_progress())
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": str(e)}
        )

@app.get("/download/merge/{request_id}")
async def download_merge(request_id: str):
    try:
        file_path = merge_paths.get(request_id)
        if not file_path:
            raise HTTPException(status_code=404, detail="文件不存在或已过期")
            
        return FileResponse(
            file_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=os.path.basename(file_path)
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download")
async def download_page(request: Request):
    return templates.TemplateResponse("download.html", {"request": request})

@app.post("/download/folder")
async def download_folder(company_folder: str = Form(...)):
    try:
        current_dir = os.getcwd()
        company_path = os.path.join(current_dir, company_folder)
        
        # 检查公司文件夹是否存在
        if not os.path.exists(company_path):
            raise HTTPException(
                status_code=404,
                detail=f"找不到公司文件夹：{company_folder}"
            )
        
        # 创建临时文件来存储 ZIP
        with tempfile.NamedTemporaryFile(delete=False, suffix='.zip') as tmp_file:
            zip_path = tmp_file.name
            
        # 创建 ZIP 文件
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(company_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    arc_name = os.path.relpath(file_path, os.path.dirname(company_path))
                    zipf.write(file_path, arc_name)
        
        # 设置下载文件名
        filename = f"{company_folder}.zip"
        headers = {
            'Content-Disposition': f'attachment; filename="{filename}"'
        }
        
        # 返回 ZIP 文件
        return FileResponse(
            zip_path,
            media_type='application/zip',
            headers=headers,
            background=BackgroundTask(lambda: os.unlink(zip_path))
        )
        
    except Exception as e:
        if 'zip_path' in locals():
            os.unlink(zip_path)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/folders")
async def get_folders():
    try:
        current_dir = os.getcwd()
        # 获取所有文件夹
        folders = [
            d for d in os.listdir(current_dir) 
            if os.path.isdir(os.path.join(current_dir, d)) 
            and len(d) >= 8 
            and d[-8:].isdigit()  # 检查后8位是否为数字（日期）
        ]
        folders.sort(reverse=True)  # 按日期倒序排序
        return folders
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000) 